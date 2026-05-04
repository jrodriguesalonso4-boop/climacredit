"""
Gerador do relatório técnico-metodológico do ClimaCredit.
Salva em docs/ClimaCredit_Relatorio_Tecnico_v1.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import copy

# ── helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=None, color=None, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    s = size or sizes.get(level, 11)
    set_font(run, size=s, bold=True, color=color)
    return p

def add_body(doc, text, justify=True, size=11, space_before=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_placeholder(doc, text):
    """Texto de aviso metodológico em itálico cinza."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"[NOTA METODOLÓGICA] {text}")
    set_font(run, size=10, italic=True, color=(150, 100, 0))
    return p

def add_verify(doc, text):
    """Marcador [VERIFICAR REFERÊNCIA EXATA]."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[VERIFICAR REFERÊNCIA EXATA] {text}")
    set_font(run, size=10, italic=True, color=(180, 0, 0))
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell_text, cell) in enumerate(zip(cells, row.cells)):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        set_font(run, size=10, bold=(bold_first and i == 0))
        p.paragraph_format.space_after = Pt(1)
    return row

def set_table_header(row, cells):
    for cell_text, cell in zip(cells, row.cells):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        set_font(run, size=10, bold=True)
        # fundo cinza claro
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E8E8E8")
        tcPr.append(shd)

def add_page_number(doc):
    """Adiciona numeração de página no rodapé."""
    section = doc.sections[0]
    footer  = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_font(run, size=9)
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"; instrText.set(qn("xml:space"), "preserve")
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    run2 = p.add_run(" / ")
    set_font(run2, size=9)
    run3 = p.add_run()
    set_font(run3, size=9)
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText"); instrText2.text = "NUMPAGES"; instrText2.set(qn("xml:space"), "preserve")
    fldChar4 = OxmlElement("w:fldChar"); fldChar4.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar3); run3._r.append(instrText2); run3._r.append(fldChar4)

def set_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
set_margins(doc)
add_page_number(doc)

hoje = date.today().strftime("%d/%m/%Y")

# ── CAPA ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("ClimaCredit")
set_font(run, size=28, bold=True, color=(0, 80, 60))
p_title.paragraph_format.space_after = Pt(6)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_sub.add_run("Monitor de Risco Climático para Crédito Agro")
set_font(run, size=18, bold=False, color=(60, 60, 60))

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_sub2.add_run("Documento Técnico-Metodológico")
set_font(run, size=13, italic=True, color=(100, 100, 100))
p_sub2.paragraph_format.space_after = Pt(30)

doc.add_paragraph()
doc.add_paragraph()

for line in [
    ("GAS Challenge 2026.1  ·  XP Inc.", 12, False),
    ("", 10, False),
    ("Equipe:", 11, True),
    ("[Preencher nomes da equipe]", 11, False),
    ("", 10, False),
    ("Mentor:", 11, True),
    ("João Gabriel Amarante", 11, False),
    ("Head de Estudos Energéticos — XP Global Markets", 10, True),
    ("", 10, False),
    (f"Data: {hoje}", 11, False),
    ("Versão: 1.0", 11, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line[0])
    set_font(run, size=line[1], bold=line[2])
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── SUMÁRIO (manual) ──────────────────────────────────────────────────────────
add_heading(doc, "Sumário", level=1, size=14)
sumario = [
    ("Sumário Executivo", "3"),
    ("1. Contexto e Problema", "4"),
    ("   1.1 O agronegócio na carteira da XP", "4"),
    ("   1.2 Eventos climáticos como driver de inadimplência agro", "4"),
    ("   1.3 Gap atual e confirmação do mentor", "4"),
    ("   1.4 Por que importa agora", "5"),
    ("2. Arquitetura da Ferramenta", "5"),
    ("   2.1 As cinco telas e o que entregam ao gestor", "5"),
    ("   2.2 Stack técnico", "5"),
    ("   2.3 Fluxo de dados", "6"),
    ("   2.4 Fontes de dados", "6"),
    ("3. Metodologia do ClimaRisk Score", "7"),
    ("   3.1 Visão geral do score composto", "7"),
    ("   3.2 Componente Precipitação (peso 35%)", "7"),
    ("   3.3 Componente ENSO (peso 25%)", "8"),
    ("   3.4 Componente Queimadas (peso 25%)", "9"),
    ("   3.5 Componente Temperatura (peso 15%)", "9"),
    ("   3.6 Score composto e thresholds de classificação", "10"),
    ("   3.7 Granularidade municipal on-demand", "10"),
    ("4. Metodologia do Tradutor Financeiro", "11"),
    ("   4.1 Crédito Agro (CRA/LCA)", "11"),
    ("   4.2 Equity Agro", "12"),
    ("   4.3 Commodities", "12"),
    ("5. Painel de Alertas", "13"),
    ("6. Calendário Agrícola × Risco", "13"),
    ("7. Limitações Metodológicas e Roadmap", "14"),
    ("8. Aderência Regulatória", "15"),
    ("9. Conclusão", "15"),
    ("Anexos", "16"),
]

tbl_sum = doc.add_table(rows=0, cols=2)
tbl_sum.style = "Table Grid"
for secao, pg in sumario:
    row = tbl_sum.add_row()
    row.cells[0].text = secao
    row.cells[1].text = pg
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=10)
            para.paragraph_format.space_after = Pt(1)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # remove borda da tabela (deixa mais limpo)
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SUMÁRIO EXECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Sumário Executivo", level=1)

add_body(doc, "O setor agropecuário representa parcela relevante das carteiras de crédito estruturado no Brasil. Certificados de Recebíveis do Agronegócio (CRA) e Letras de Crédito do Agronegócio (LCA) estão entre os instrumentos de maior crescimento no mercado de capitais brasileiro, impulsionados pelo FIAGRO e pela demanda institucional. Essa expansão traz exposição climática implícita que, até hoje, não dispõe de ferramenta integrada e de fácil operação para os gestores de portfólio.")

add_body(doc, "O ClimaCredit nasce dessa lacuna. O sistema agrega dados climáticos públicos de fontes primárias brasileiras e internacionais — precipitação, ENSO, focos de queimada e temperatura — e os traduz em linguagem financeira acionável: probabilidade de default ajustada ao clima (PD), perda esperada (EL), alertas por perfil de investidor e monitoramento de empresas listadas.")

add_body(doc, "O diferencial central da ferramenta é a rastreabilidade: cada número do dashboard é derivável até sua fonte primária. Quando uma calibração é preliminar, o sistema documenta isso explicitamente — honestidade metodológica é premissa de projeto, não exceção.")

add_body(doc, "Resultados do MVP atual: cobertura nacional (27 UFs + busca municipal sob demanda), 7 culturas no calendário agrícola, score composto com 4 componentes ponderadas, 5 telas funcionais cobrindo dashboard, calendário, ranking por UF, alertas diferenciados por perfil e tradutor financeiro. O sistema roda inteiramente sobre APIs públicas, sem dependência de dados proprietários.")

add_body(doc, "Do ponto de vista regulatório, o ClimaCredit está alinhado com a BCB Resolução 139 (2021) [VERIFICAR REFERÊNCIA EXATA — confirmar número e data da resolução], que exige que instituições financeiras identifiquem, avaliem e monitorem riscos climáticos em suas operações de crédito. A granularidade espacial e a decomposição metodológica transparente da ferramenta respondem diretamente a esse requisito.")

add_verify(doc, "Confirmar número e ano exato da BCB Resolução sobre risco climático em crédito antes de enviar o relatório.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. CONTEXTO E PROBLEMA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Contexto e Problema", level=1)

add_heading(doc, "1.1  O agronegócio na carteira da XP", level=2)
add_body(doc, "O agronegócio responde por aproximadamente 25% do PIB brasileiro e é o setor de maior peso nos instrumentos de crédito privado negociados no mercado de capitais nacional. CRA e LCA apresentaram crescimento acumulado superior a 200% entre 2019 e 2024 [VERIFICAR REFERÊNCIA EXATA — ANBIMA Boletim de Mercado de Capitais], impulsionados pela demanda por alternativas ao crédito bancário e pelos incentivos tributários para pessoa física. Para a XP, como plataforma líder em distribuição de CRA e gestão de FIAGRO, a exposição climática implícita nessas carteiras é material e crescente.")

add_body(doc, "Essa exposição, no entanto, não se traduz facilmente em métricas financeiras. Um gestor que monitora um CRA com lastro em soja do Mato Grosso precisaria, hoje, cruzar manualmente dados do INPE, CONAB, NOAA e Open-Meteo para avaliar se a anomalia climática da safra atual justifica revisão de provisão — processo que o ClimaCredit automatiza em tempo real.")

add_heading(doc, "1.2  Eventos climáticos como driver de inadimplência agro", level=2)
add_body(doc, "A correlação entre eventos climáticos extremos e inadimplência no crédito rural brasileiro está documentada na literatura. O Banco Central do Brasil, em suas Notas de Crédito Rural anuais, registra elevação consistente de inadimplência nos anos de El Niño severo (2015-16) e La Niña prolongada (2020-22) nas regiões mais sensíveis. O BNDES e o Banco Mundial publicaram estudos relacionando déficit hídrico e frustração de safra com risco de crédito em operações de custeio e investimento rural [VERIFICAR REFERÊNCIA EXATA — identificar papers específicos BCB/BNDES sobre inadimplência agro e clima].")

add_body(doc, "Eventos recentes reforçam esse vínculo: a seca histórica de 2021-2022 no Sul do Brasil elevou a sinistralidade do PROAGRO a níveis recordes; o El Niño de 2023-2024 comprometeu a segunda safra de milho no Centro-Oeste; e as queimadas de agosto-setembro de 2024 geraram impacto em garantias produtivas de operações CPR em MG, GO e MT.")

add_heading(doc, "1.3  Gap atual e confirmação do mentor", level=2)
add_body(doc, "João Gabriel Amarante, Head de Estudos Energéticos da XP Global Markets e mentor do projeto, confirmou a ausência de ferramenta integrada de monitoramento de risco climático físico para as carteiras de CRA/LCA da XP. As análises existentes são pontuais, manuais e não integradas ao fluxo de decisão de portfólio. O ClimaCredit é a resposta direta a esse gap identificado pelo próprio ecossistema XP.")

add_heading(doc, "1.4  Por que importa agora", level=2)
add_body(doc, "Quatro forças convergem para tornar o monitoramento de risco climático urgente no contexto atual:")
add_bullet(doc, "Regulação: BCB Resolução 139 exige que instituições financeiras incorporem risco climático em seus processos de gestão de risco [VERIFICAR REFERÊNCIA EXATA]. A convergência com TCFD e NGFS aumenta a pressão por divulgação estruturada.")
add_bullet(doc, "Ciclo ENSO ativo: o sistema ENSO saiu de El Niño moderado (2023-24) e está em transição. O impacto sobre padrões de chuva no Brasil é imediato e geograficamente heterogêneo — exatamente o que o ClimaCredit monitora.")
add_bullet(doc, "Intensificação de eventos extremos: a frequência e magnitude de secas, chuvas extremas e queimadas no Brasil aumentou consistentemente na última década, com impacto documentado em produtividade agrícola.")
add_bullet(doc, "Crescimento dos FIAGRO: fundos que investem diretamente em CRA e CPR precisam de ferramentas de due diligence climática que ainda não existem no mercado.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARQUITETURA DA FERRAMENTA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Arquitetura da Ferramenta", level=1)

add_heading(doc, "2.1  As cinco telas e o que entregam ao gestor", level=2)

tbl_telas = doc.add_table(rows=0, cols=3)
tbl_telas.style = "Table Grid"
set_table_header(tbl_telas.add_row(), ["Tela", "Nome", "Entrega ao gestor"])
rows_telas = [
    ("1", "Dashboard Principal", "Score ClimaRisk por região (0-100), decomposição em 4 componentes, histórico ONI e busca municipal on-demand"),
    ("2", "Calendário Agrícola × Risco", "Cruzamento cultura × região × fase atual do ciclo (plantio/desenvolvimento/colheita) com score de risco sobreposto"),
    ("3", "ClimaRisk Score por UF", "Ranking dos 27 estados com score, mapa coroplético interativo e série histórica por UF"),
    ("4", "Painel de Alertas", "Alertas ativos por perfil: Crédito Agro (PD/provisão), Equity Agro (EBITDA/ticker), Commodities (hedge/curva futura)"),
    ("5", "Tradutor Financeiro", "Simulador CRA/LCA com PD ajustada ao clima, banco de empresas listadas com score climático ponderado, e impacto ENSO × preço de commodities"),
]
for row_data in rows_telas:
    add_table_row(tbl_telas, row_data, bold_first=True)

doc.add_paragraph()

add_heading(doc, "2.2  Stack técnico", level=2)
add_body(doc, "A ferramenta é construída inteiramente em Python. O processamento de dados ocorre em scripts independentes que geram arquivos CSV intermediários. O dashboard é servido via Streamlit, framework de prototipagem rápida adequado para MVP. A escolha por Streamlit é consciente e explicitamente temporária: para um produto de produção, a migração para React/FastAPI seria o passo natural, mas está fora do escopo do challenge.")

add_bullet(doc, "ETL e processamento: Python 3.11 — scripts em scripts/")
add_bullet(doc, "Dashboard: Streamlit 1.x + Plotly Express/Graph Objects")
add_bullet(doc, "Manipulação de dados: pandas, numpy")
add_bullet(doc, "Visualização de mapas: Folium + streamlit-folium, Plotly choropleth")
add_bullet(doc, "Geocodificação e vizinho mais próximo: scikit-learn BallTree com métrica haversine")
add_bullet(doc, "HTTP: requests, urllib3")

add_heading(doc, "2.3  Fluxo de dados", level=2)
add_body(doc, "O pipeline segue uma arquitetura de camadas sequenciais:")
add_body(doc, "FONTES EXTERNAS → SCRIPTS DE ETL → CSVs INTERMEDIÁRIOS → APP.PY (leitura + cálculo on-demand) → DASHBOARD")
add_body(doc, "Cada script é independente e idempotente: pode ser re-executado a qualquer momento sem efeito colateral. O botão '↺ Atualizar' no dashboard dispara os quatro scripts em sequência (ENSO/NOAA → Queimadas/FIRMS → Clima/ERA5 → Score). O tempo total de atualização é de 60 a 120 segundos, dominado pelas 27 chamadas à API Open-Meteo (uma por UF, com delay de 0,3s para respeitar rate limits).")
add_body(doc, "Buscas municipais on-demand usam @st.cache_data com TTL de 24 horas. A justificativa: dados climáticos via ERA5 não mudam de hora em hora (ERA5 é reanalysis com delay de 5 dias), e o cache evita latência de 3-5 segundos a cada digitação no campo de busca.")

add_heading(doc, "2.4  Fontes de dados", level=2)

tbl_fontes = doc.add_table(rows=0, cols=5)
tbl_fontes.style = "Table Grid"
set_table_header(tbl_fontes.add_row(), ["Fonte", "Institução", "Dado", "Frequência", "Granularidade / Limitações"])
rows_fontes = [
    ("NOAA PSL\nnina34.anom.data", "NOAA (EUA)", "Índice ONI\n(Oceanic Niño Index)", "Mensal\n(delay ~30d)", "Nacional / Global — sem granularidade regional direta"),
    ("NASA FIRMS\nVIIRS SNPP", "NASA (EUA)", "Focos de calor\n(últimos 7 dias)", "Diária", "Ponto geográfico (lat/lon) — cobertura global. Falha ocasional no endpoint do INPE motivou uso do FIRMS como fonte primária"),
    ("Open-Meteo\nArchive API (ERA5)", "Open-Meteo /\nCopernicus/ECMWF", "Precipitação e\ntemperatura diária", "Diária\n(delay ~5 dias ERA5)", "Grade 0,25° × 0,25° (~28km) — sem autenticação. Limitação: ERA5 é reanalysis, não observação direta"),
    ("IBGE", "IBGE (BR)", "Geocódigos municipais,\ncoordenadas, censo agropecuário 2017", "Censo/estático", "Municipal — área agrícola usada para ponderação de focos"),
    ("INMET\n(Normais 1991-2020)", "INMET (BR)", "Normais climatológicas\nde precipitação e temperatura", "Estático\n(publicação 2022)", "Estadual (capital) — precisão ±1-2°C adequada para anomalias relativas"),
    ("CONAB\n(publicação oficial)", "CONAB (BR)", "Calendário agrícola:\nplantio, desenvolvimento, colheita", "Estático\n(sem API direta)", "Regional — montado manualmente a partir da publicação oficial. Atualização anual necessária"),
]
for row_data in rows_fontes:
    add_table_row(tbl_fontes, row_data)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. METODOLOGIA DO CLIMARISK SCORE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Metodologia do ClimaRisk Score", level=1)

add_heading(doc, "3.1  Visão geral do score composto", level=2)
add_body(doc, "O ClimaRisk Score é um índice composto de 0 a 100 que agrega quatro componentes climáticos em uma única métrica de risco agrícola. Quanto maior o score, maior o risco de impacto climático sobre a produção. A arquitetura em componentes separadas serve a dois propósitos: (1) permitir que o gestor decomponha o risco — um score de 75 por queimada é qualitativamente diferente de um score de 75 por seca; e (2) facilitar a atualização metodológica de cada componente de forma independente.")

add_body(doc, "A ferramenta calcula o score em duas granularidades:")
add_bullet(doc, "Regional (5 macrorregiões): calculado em scripts/score.py — serve o dashboard principal e o tradutor financeiro")
add_bullet(doc, "Estadual (27 UFs): calculado em scripts/score_ufs.py — serve o mapa e ranking por UF")

add_body(doc, "O score final é a soma ponderada das quatro componentes normalizadas individualmente para o intervalo [0, 100]:")

p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_formula.add_run("Score = 0,35 × S_precip + 0,25 × S_ENSO + 0,25 × S_queimadas + 0,15 × S_temp")
set_font(run, size=11, bold=True)

add_heading(doc, "3.2  Componente Precipitação (peso 35%)", level=2)
add_body(doc, "Fonte: Open-Meteo Archive API (ERA5 reanalysis), endpoint https://archive-api.open-meteo.com/v1/archive. Variável: precipitation_sum diária. Janela: 30 dias anteriores ao dia atual.")
add_body(doc, "A precipitação acumulada observada é comparada à normal climatológica 1991-2020 do INMET para o mês corrente, ajustada proporcionalmente pela janela de 30 dias. A anomalia percentual é calculada como:")

p_f1 = doc.add_paragraph()
p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_f1.add_run("anomalia_pct = (P_obs − P_normal) / max(P_normal, 1) × 100")
set_font(run, size=11, italic=True)

add_body(doc, "A função de normalização é assimétrica: tanto seca severa quanto excesso hídrico elevam o score, mas com inclinações diferentes, porque o impacto agrícola de cada extremo é qualitativamente distinto. A função por faixas implementada é:")

tbl_prec = doc.add_table(rows=0, cols=3)
tbl_prec.style = "Table Grid"
set_table_header(tbl_prec.add_row(), ["Faixa de anomalia", "Fórmula S_precip", "Interpretação"])
rows_prec = [
    ("anomalia ≤ −50%", "min(100, 90 + |a+50| × 0,2)", "Seca severa a extrema"),
    ("−50% < anomalia < −20%", "40 + (|a|−20) × (50/30)", "Seca moderada"),
    ("−20% ≤ anomalia ≤ +20%", "max(0, 10 + |a| × 1,5)", "Condição normal (score baixo)"),
    ("+20% < anomalia ≤ +80%", "20 + (a−20) × (40/60)", "Excesso moderado"),
    ("anomalia > +80%", "min(100, 60 + (a−80) × 0,3)", "Excesso severo"),
]
for r in rows_prec:
    add_table_row(tbl_prec, r)
doc.add_paragraph()

add_body(doc, "Justificativa do peso 35%: a precipitação é o driver climático de impacto mais direto e imediato sobre a produtividade agrícola, com correlação mais robusta e menos defasada do que temperatura ou ENSO. A revisão bibliográfica disponível (EMBRAPA, CONAB, literatura agroclimatológica) consistentemente aponta déficit hídrico como o principal fator de frustração de safra no Brasil [VERIFICAR REFERÊNCIA EXATA — citar paper ou relatório específico EMBRAPA/CONAB sobre impacto de precipitação em produtividade].")

add_heading(doc, "3.3  Componente ENSO (peso 25%)", level=2)
add_body(doc, "Fonte: NOAA Physical Sciences Laboratory — índice ONI (Oceanic Niño Index), calculado como anomalia de temperatura da superfície do mar na região Niño 3.4. Valores positivos indicam El Niño; negativos, La Niña. Threshold convencional: |ONI| ≥ 0,5°C por 5 meses consecutivos.")
add_body(doc, "O score ENSO é calculado como:")

p_f2 = doc.add_paragraph()
p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_f2.add_run("S_ENSO = max(0, min(100, (ONI × sens_UF + 1,5) / 3,0 × 100))")
set_font(run, size=11, italic=True)

add_body(doc, "Onde sens_UF é o fator de sensibilidade ENSO da UF, que varia de −0,85 (Sul — La Niña aumenta risco) a +0,95 (CE — El Niño aumenta risco). O denominador 3,0 normaliza o produto para o intervalo [0,1], assumindo que ONI = ±3,0°C representa evento extremo. O valor 1,5 é o offset que garante que ONI = 0 (condição neutra) resulte em S_ENSO = 50.")
add_body(doc, "Os fatores de sensibilidade são documentados em scripts/score_ufs.py com referência às fontes primárias. A tabela completa por UF está no Anexo A. As faixas regionais principais são:")

tbl_enso = doc.add_table(rows=0, cols=3)
tbl_enso.style = "Table Grid"
set_table_header(tbl_enso.add_row(), ["Região / UFs", "Sensibilidade (sens)", "Base científica"])
rows_enso = [
    ("Nordeste semiárido (CE, RN, PB, PE)", "+0,85 a +0,95", "El Niño → bloqueio da ZCIT → seca severa. Correlação histórica robusta. Grimm & Tedeschi (2009); Ropelewski & Halpert (1987) [VERIFICAR]"),
    ("Sul (RS, SC, PR)", "−0,70 a −0,85", "La Niña → bloqueio da frente polar → seca primavera-verão. Grimm et al. (2000) [VERIFICAR]"),
    ("Norte/Amazônia (AM, PA, AC, RO, TO)", "+0,45 a +0,65", "El Niño → seca amazônica, aumento de focos. Marengo et al. (2008) [VERIFICAR]"),
    ("Centro-Oeste (MT, MS, GO, DF)", "+0,55 a +0,65", "El Niño → seca do Cerrado, atraso no início das chuvas"),
    ("Sudeste (SP, MG, RJ, ES)", "+0,25 a +0,40", "Efeito moderado e geograficamente heterogêneo"),
    ("Maranhão (MA)", "+0,65", "Transição NE/Amazônia; predominantemente El Niño → seca"),
]
for r in rows_enso:
    add_table_row(tbl_enso, r)
doc.add_paragraph()

add_placeholder(doc, "Os fatores de sensibilidade foram calibrados com base na literatura científica disponível e no Atlas Climático do CPTEC/INPE. Para V2, recomenda-se calibração empírica a partir de correlação estatística entre série histórica de ONI e séries de produtividade/safra do IBGE PAM por UF.")

add_heading(doc, "3.4  Componente Queimadas (peso 25%)", level=2)
add_body(doc, "Fonte: NASA FIRMS VIIRS SNPP — arquivo de focos de calor dos últimos 7 dias, endpoint CSV sem autenticação. A fonte primária INPE BDQueimadas foi testada mas apresentou instabilidade de endpoint durante o desenvolvimento, motivando a adoção do NASA FIRMS como fonte principal.")
add_body(doc, "Na granularidade regional (5 regiões), cada foco é associado à sua região a partir do estado de origem (coluna estado). Os focos são normalizados pelo P95 histórico regional, estimado a partir do arquivo histórico NASA FIRMS VIIRS SNPP para o Brasil (2012-2023):")

p_f3 = doc.add_paragraph()
p_f3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_f3.add_run("S_queimadas = min(100, focos_7d / P95_regional × 100)")
set_font(run, size=11, italic=True)

add_body(doc, "Na granularidade estadual (27 UFs), a normalização usa percentile rank relativo entre as 27 UFs no período atual, com método min para empates. Essa abordagem foi adotada porque a distribuição de focos por UF é fortemente assimétrica (log-normal), e qualquer threshold absoluto calibrado para o pico de seca (julho-outubro) trivialmente satura em meses de transição. O percentile rank garante poder discriminatório permanente independente da época do ano.")

add_placeholder(doc, "Os valores de P95 regional (Centro-Oeste: 6.000, Norte: 6.000, Sul: 3.000, Sudeste: 3.500, Nordeste: 2.500 focos/7 dias) são estimativas baseadas no arquivo histórico NASA FIRMS. Para V2, recomenda-se cálculo formal de P95 sobre série histórica completa por região e por mês, para controlar sazonalidade.")

add_heading(doc, "3.5  Componente Temperatura (peso 15%)", level=2)
add_body(doc, "Fonte: Open-Meteo Archive API (ERA5) — variáveis temperature_2m_max e temperature_2m_min. Temperatura média diária = (Tmax + Tmin) / 2. Anomalia = temperatura média observada (30d) − normal climatológica INMET 1991-2020 para o mês corrente.")
add_body(doc, "A função de normalização trata o Sul de forma distinta:")
add_bullet(doc, "Sul (RS, SC, PR) com anomalia < −1,5°C: score = min(100, |anomalia| × 22). Geadas severas (−4,5°C) → score próximo a 100.")
add_bullet(doc, "Demais UFs/regiões: score = min(100, max(0, anomalia × 14 + 5)). Calibração: 0°C → score 5 (sem risco), +3°C → score 47 (atenção), +5°C → score 75 (crítico), +6,8°C → satura.")

add_body(doc, "O slope 14 foi escolhido porque o slope anterior (25) saturava a +3,2°C — anomalias de +4-5°C são incomuns mas ocorrem (ex: RS e SC em abril/2026). Com slope 14, esses eventos ficam em 63-73, preservando discriminação para eventos extremamente severos.")

add_body(doc, "INCONSISTÊNCIA DOCUMENTADA: Na granularidade regional (scripts/score.py), a componente temperatura não usa dados ERA5 — usa S_ENSO × 0,7 como proxy, por limitação técnica do script regional. Na granularidade estadual (scripts/score_ufs.py), a componente temperatura usa ERA5 de verdade. Essa inconsistência é conhecida e está na lista de correções para V2.")

add_placeholder(doc, "A componente temperatura regional é proxy do ENSO, não dado real. O peso de 15% foi mantido para consistência de fórmula entre as granularidades, mas o score regional deve ser interpretado com essa limitação em mente. A substituição por ERA5 real na granularidade regional é prevista para V2.")

add_heading(doc, "3.6  Score composto e thresholds de classificação", level=2)
add_body(doc, "O score final é a soma ponderada dos quatro componentes normalizados. Os thresholds de classificação são:")

tbl_thresh = doc.add_table(rows=0, cols=3)
tbl_thresh.style = "Table Grid"
set_table_header(tbl_thresh.add_row(), ["Faixa", "Classificação", "Interpretação operacional"])
rows_thresh = [
    ("score < 45", "NORMAL", "Condições climáticas sem alerta — monitoramento padrão"),
    ("45 ≤ score < 70", "ATENÇÃO", "Anomalia climática em desenvolvimento — revisão de exposição recomendada"),
    ("score ≥ 70", "CRÍTICO", "Evento climático severo ativo — ação imediata de gestão de risco"),
]
for r in rows_thresh:
    add_table_row(tbl_thresh, r)
doc.add_paragraph()

add_placeholder(doc, "Os thresholds 45 e 70 foram definidos por calibração intuitiva baseada na distribuição esperada de scores: aproximadamente 60% NORMAL, 30% ATENÇÃO, 10% CRÍTICO em condições médias. Para V2, recomenda-se backtesting do score em eventos históricos documentados (El Niño 2015-16, La Niña 2020-22, seca sul 2021-22) para validação e eventual recalibração dos thresholds.")

add_heading(doc, "3.7  Granularidade municipal on-demand", level=2)
add_body(doc, "O gestor pode buscar qualquer município brasileiro pelo nome e obter um card com o score local, decomposição em componentes e fase do ciclo agrícola das culturas predominantes na microrregião. O cálculo é realizado em tempo real via Open-Meteo/ERA5 para a coordenada do município.")
add_body(doc, "Para municípios sem coordenadas mapeadas diretamente no banco interno (aproximadamente 200 municípios agrícolas cobertos), o sistema usa busca por vizinho mais próximo via BallTree com métrica Haversine (scikit-learn). O vizinho mais próximo é identificado entre os municípios do banco, e o gestor recebe um aviso explícito indicando a distância de fallback e o município utilizado como referência.")
add_body(doc, "O cache de 24 horas (@st.cache_data com TTL=86400s) garante que buscas repetidas do mesmo município não incorram em latência adicional. A justificativa técnica: ERA5 opera com delay de 5 dias e é atualizado diariamente; variações climáticas relevantes para risco agro ocorrem em escala de dias, não horas.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. METODOLOGIA DO TRADUTOR FINANCEIRO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Metodologia do Tradutor Financeiro", level=1)

add_heading(doc, "4.1  Crédito Agro (CRA/LCA)", level=2)
add_body(doc, "O módulo de Crédito Agro é um simulador de carteira que aceita um conjunto de operações (região, cultura, valor exposto em R$ mi, prazo em anos) e calcula métricas de risco de crédito ajustadas ao clima para cada posição e para a carteira total.")

add_body(doc, "ESTADO ATUAL DA IMPLEMENTAÇÃO (V1):")
add_bullet(doc, "PD base: 3,0%/ano — calibração preliminar baseada em ordem de grandeza do setor agro. Fonte declarada no dashboard: 'histórico setor agro'. Não há citação específica a série do BCB/SCR na V1.")
add_bullet(doc, "Multiplicadores climáticos: NORMAL ×1,0 / ATENÇÃO ×1,4 / CRÍTICO ×2,2. Calibração arbitrária — sem referência empírica documentada.")
add_bullet(doc, "LGD: 100% implícito — a perda esperada é calculada como EAD × PD_acumulada, sem fator de recuperação.")
add_bullet(doc, "Fórmula PD acumulada: 1 − (1 − PD_ano)^prazo — modelo de hazard rate composto, matematicamente correto.")
add_bullet(doc, "Concentração: pizza descritiva por região, sem métrica HHI ou índice de concentração calculado.")

add_placeholder(doc, "A V1 do Crédito Agro é um framework funcional com calibração preliminar. Os multiplicadores ×1,4 e ×2,2 e a PD base de 3,0% precisam de ancoragem empírica antes da apresentação final. Para a V2, a proposta é: (1) ancorar PD base na série histórica de inadimplência rural do BCB SCR 2019-2023; (2) calibrar multiplicadores a partir de correlação entre ONI e inadimplência agro por região; (3) adicionar LGD = 50% (faixa típica para crédito rural com garantia real, alinhado com Basel III); (4) calcular EL = EAD × PD_acumulada × LGD; (5) implementar HHI de concentração regional.")

add_body(doc, "METODOLOGIA-ALVO (V2) — documentada para orientar a continuidade do desenvolvimento:")
add_body(doc, "PD_ano = PD_base × fator_clima × fator_cultura. EL = EAD × PD_acumulada × LGD. PD_acumulada = 1 − (1 − PD_ano)^prazo. Ajuste de concentração: stress_factor = 1 + HHI × 0,30, onde HHI = Σ(share_regiao²). LGD-alvo: 50% (referência Basel III para crédito garantido por ativo real) [VERIFICAR REFERÊNCIA EXATA].")

add_heading(doc, "4.2  Equity Agro", level=2)
add_body(doc, "O módulo Equity Agro mantém um banco de 15 empresas listadas com mapeamento de exposição geográfica por macrorregião, derivado de Relatórios Anuais e ITRs. O score climático de cada empresa é calculado como a média ponderada do ClimaRisk Score das regiões onde opera:")

p_f4 = doc.add_paragraph()
p_f4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_f4.add_run("Score_empresa = Σ (exp_regiao_i × Score_regiao_i)")
set_font(run, size=11, italic=True)

add_body(doc, "As empresas cobertas incluem SLCE3, AGRO3, TTEN3, SMTO3, CAML3, RAIZ4, BEEF3, JBSS3, BRFS3, MRFG3, MDIA3, CSAN3, SUZB3, KLBN3 e VBBR3. Para empresas fora do banco, o sistema retorna mensagem explícita informando a limitação de cobertura.")
add_placeholder(doc, "A exposição geográfica das empresas foi mapeada manualmente a partir de documentos públicos. Atualização necessária após cada resultado trimestral ou evento corporativo relevante. A V2 prevê automação da atualização via leitura de RAA/ITR.")

add_heading(doc, "4.3  Commodities", level=2)
add_body(doc, "ESTADO ATUAL DA IMPLEMENTAÇÃO (V1):")
add_bullet(doc, "Três commodities cobertas: Soja, Milho, Café. Expansão para Arroz, Algodão e Açúcar está prevista para V2.")
add_bullet(doc, "Impacto esperado: ip = ONI_atual × coeficiente × 100. Coeficientes: Soja = 0,06; Milho = 0,04; Café = 0,08.")
add_bullet(doc, "Série histórica no gráfico: ONI × coeficiente para os últimos 60 meses — é proxy reescalado do ONI, não série real de preço.")
add_bullet(doc, "Câmbio USD/BRL, estoques globais (USDA WASDE), sazonalidade e prêmio de base não estão modelados na V1.")
add_bullet(doc, "Intervalo de confiança: não implementado — ponto único sem faixa de incerteza.")

add_placeholder(doc, "A referência Iizumi et al. (2014) documentada no dashboard mede impacto em yield (toneladas/hectare), não diretamente em preço. A tradução de variação de oferta para variação de preço exige uma elasticidade, que não está implementada na V1. O texto do dashboard deve ser interpretado como proxy de ordem de grandeza, não como previsão de preço. A V2 incorporará: (1) conexão com série histórica CEPEA/ESALQ para calibração empírica; (2) elasticidade preço-oferta por commodity; (3) incorporação de câmbio USD/BRL; (4) intervalo de confiança via Monte Carlo sobre distribuição histórica de ONI.")
add_verify(doc, "Confirmar referência exata: Iizumi, T. et al. (2014). 'Impacts of El Niño Southern Oscillation on the global yields of major crops.' Nature Communications. Citar volume, número e DOI.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PAINEL DE ALERTAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Painel de Alertas", level=1)

add_body(doc, "O Painel de Alertas consolida os sinais climáticos ativos em cartões acionáveis, diferenciados por perfil de investidor. A lógica de geração aplica thresholds sobre os dados climáticos mais recentes:")

tbl_alertas = doc.add_table(rows=0, cols=3)
tbl_alertas.style = "Table Grid"
set_table_header(tbl_alertas.add_row(), ["Alerta", "Condição de disparo", "Perfis que recebem"])
rows_alertas = [
    ("Déficit hídrico", "anomalia_precip < −30%", "Crédito Agro / Equity Agro / Commodities"),
    ("Excesso hídrico", "anomalia_precip > +40%", "Crédito Agro / Commodities"),
    ("Focos de queimada elevados", "focos_7d > 200", "Crédito Agro / Equity Agro"),
    ("ENSO ativo", "|ONI| ≥ 0,5°C", "Crédito Agro / Equity Agro / Commodities"),
]
for r in rows_alertas:
    add_table_row(tbl_alertas, r)
doc.add_paragraph()

add_body(doc, "Cada alerta possui campos de impacto e ação diferenciados por perfil. A diferenciação não é apenas cosmética — os campos são escritos em linguagem específica de cada audiência:")
add_bullet(doc, "Crédito Agro: fala a língua de PD ajustada, provisão IFRS9, colateral e CPR físico. Cita a variação de PD estimada com o multiplicador climático aplicado.")
add_bullet(doc, "Equity Agro: fala a língua de EBITDA, guidance e tickers. Cita os tickers do EQUITY_DB com maior exposição à região alertada, com percentual de exposição.")
add_bullet(doc, "Commodities: fala a língua de hedge, curva futura e basis local. Cita os contratos relevantes na B3/CBOT para a região e commodity afetadas.")
add_body(doc, "Os alertas são ordenados por score (maior para menor) dentro de cada perfil. Alertas ativos são preservados entre atualizações dos dados climáticos.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. CALENDÁRIO AGRÍCOLA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Calendário Agrícola × Risco", level=1)

add_body(doc, "A Tela 2 cruza o calendário de ciclo das culturas com o score de risco climático vigente, respondendo à pergunta operacional: 'dado o risco climático atual, qual fase do ciclo está mais exposta e onde?'")
add_body(doc, "Cobertura: 7 culturas × 5 regiões = 35 combinações. As 7 culturas são: Soja, Milho 1ª safra, Cana, Café, Algodão, Arroz e Feijão. As datas de plantio, desenvolvimento e colheita foram extraídas da publicação oficial do CONAB (Calendário de Plantio e Colheita) e complementadas com informações do IBGE PAM (Produção Agrícola Municipal).")
add_body(doc, "O heatmap apresenta três estados visuais:")
add_bullet(doc, "Score por cor (verde → amarelo → vermelho): 27 combinações com produção em escala comercial")
add_bullet(doc, "⊘ cobertura baixa: 2 combinações onde existe produção mas é marginal (Arroz × Sudeste — Vale do Ribeira/SP; Café × Centro-Oeste — polo emergente em GO/Cristalina). Fonte: IBGE PAM 2022.")
add_bullet(doc, "N/D não aplicável (cinza hachurado): 6 combinações onde a cultura não é cultivada em escala comercial (Cana × Norte, Café × Sul, Algodão × Sul/Sudeste/Norte, Feijão × Norte)")

# ══════════════════════════════════════════════════════════════════════════════
# 7. LIMITAÇÕES E ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "7. Limitações Metodológicas e Roadmap", level=1)

add_body(doc, "Esta seção é parte central da honestidade metodológica do projeto. Listar limitações abertamente não enfraquece a proposta — demonstra rigor e permite que a audiência avalie o que é sólido do que é preliminar.")

add_heading(doc, "7.1  Calibrações preliminares que precisam de validação empírica", level=2)
add_bullet(doc, "Multiplicadores de PD climática (×1,4 / ×2,2): calibração de ordem de grandeza. Precisam de correlação estatística entre série histórica de ONI/anomalia e inadimplência rural por região (BCB SCR ou BACEN SCR).")
add_bullet(doc, "PD base 3,0%/ano: valor típico do setor mas sem citação específica a fonte primária. Ancoragem recomendada: BCB Nota de Crédito Rural — inadimplência do crédito rural por modalidade (2019-2023).")
add_bullet(doc, "Sensibilidade ENSO por UF: baseada em literatura científica mas não calibrada empiricamente para inadimplência agro. A literatura disponível correlaciona ONI com produtividade, não diretamente com default.")
add_bullet(doc, "Coeficientes ENSO × commodity (Soja: 0,06; Milho: 0,04; Café: 0,08): sem calibração documentada. A referência Iizumi et al. (2014) é para yield, não preço.")
add_bullet(doc, "Thresholds NORMAL/ATENÇÃO/CRÍTICO (45/70): definidos por intuição. Backtesting em eventos históricos pendente.")

add_heading(doc, "7.2  Componentes ainda não modeladas", level=2)
add_bullet(doc, "LGD no módulo de Crédito Agro: V1 assume LGD = 100% implicitamente. V2 incorporará LGD = 50% com referência Basel III para crédito rural garantido.")
add_bullet(doc, "Estoques globais (USDA WASDE, CONAB) no módulo de Commodities")
add_bullet(doc, "Posição na curva futura (CBOT/B3) e prêmio de base local (spread Paranaguá vs. CBOT)")
add_bullet(doc, "Taxa de câmbio USD/BRL no módulo de Commodities — relevante porque soja/milho têm preço referenciado em dólar")
add_bullet(doc, "Eventos climáticos não-ENSO: Oscilação de Madden-Julian (MJO), dipolo do Atlântico, Zona de Convergência do Atlântico Sul (ZCAS)")
add_bullet(doc, "Matriz de correlação de defaults inter-regional: a V1 trata cada operação como independente")
add_bullet(doc, "HHI de concentração regional na carteira CRA/LCA")
add_bullet(doc, "Fator cultura-específico na PD: V1 usa apenas risco regional, sem diferenciar soja de cana na mesma região")

add_heading(doc, "7.3  Gargalos de dado", level=2)
add_bullet(doc, "INPE BDQueimadas: endpoint apresentou instabilidade durante o desenvolvimento. Resolvido com fallback para NASA FIRMS como fonte primária.")
add_bullet(doc, "Open-Meteo ERA5: reanalysis com delay de ~5 dias e resolução de 28km. Para municípios pequenos em relevo complexo (serras, vales), a grade pode não representar bem as condições locais.")
add_bullet(doc, "Normais climatológicas INMET 1991-2020: usadas via capital estadual como proxy para toda a UF. Precisão ±1-2°C para temperatura; ±15-20% para precipitação em estados com heterogeneidade climática alta (BA, MT, AM).")
add_bullet(doc, "Calendário CONAB: não há API pública. Montado manualmente a partir da publicação oficial — requer atualização manual anual.")
add_bullet(doc, "Score regional (5 regiões): temperatura é proxy ENSO × 0,7, não ERA5 real. Score estadual (27 UFs) usa ERA5 de verdade.")

add_heading(doc, "7.4  Roadmap V2", level=2)
rows_road = [
    ("Curto prazo\n(pré-apresentação)", "Ancoragem de PD base em BCB SCR; LGD = 50% no simulador CRA/LCA; HHI de concentração; corrigir temperatura regional para ERA5 real"),
    ("Médio prazo\n(pós-challenge)", "Conexão com série CEPEA/ESALQ para calibração de coeficientes commodity; elasticidade preço-oferta; câmbio USD/BRL; intervalo de confiança em preços"),
    ("Longo prazo\n(produto XP)", "Integração com carteira CRA/LCA real via upload protegido; backtesting do score em eventos históricos; modelagem de correlação inter-regional; automação de relatórios para comitê de risco"),
]
tbl_road = doc.add_table(rows=0, cols=2)
tbl_road.style = "Table Grid"
set_table_header(tbl_road.add_row(), ["Horizonte", "Iniciativas"])
for r in rows_road:
    add_table_row(tbl_road, r)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. ADERÊNCIA REGULATÓRIA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Aderência Regulatória", level=1)

add_body(doc, "A BCB Resolução CMN n.º 4.557 (2017) e a Resolução BCB n.º 139 (2021) [VERIFICAR REFERÊNCIA EXATA — confirmar a resolução específica sobre risco climático] estabelecem requisitos de identificação, avaliação e monitoramento de riscos em instituições financeiras, incluindo riscos socioambientais e climáticos. O ClimaCredit endereça esses requisitos de três formas:")
add_bullet(doc, "Rastreabilidade: cada componente do score é derivável até sua fonte primária pública (NOAA, NASA, Open-Meteo/ERA5, INMET). O sistema exibe a fonte de cada dado diretamente na interface.")
add_bullet(doc, "Transparência metodológica: pesos, fórmulas e thresholds estão documentados no código e neste relatório. Calibrações preliminares são sinalizadas explicitamente.")
add_bullet(doc, "Granularidade adequada: score disponível em três níveis (regional, estadual e municipal on-demand), permitindo que a análise de risco seja calibrada ao nível de cada operação de crédito.")

add_body(doc, "Do ponto de vista de convergência internacional, o ClimaCredit está alinhado com as recomendações do TCFD (Task Force on Climate-related Financial Disclosures) para divulgação de métricas de risco climático físico, e com o framework do NGFS (Network for Greening the Financial System) para integração de risco climático em modelos de risco de crédito.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. CONCLUSÃO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Conclusão", level=1)

add_body(doc, "O ClimaCredit entrega um MVP funcional que resolve o problema identificado: transformar dados climáticos dispersos em linguagem financeira acionável para gestores de portfólio CRA/LCA. A arquitetura modular, o pipeline rastreável e a documentação aberta de limitações posicionam o projeto como base técnica sólida para evolução.")
add_body(doc, "Os próximos passos antes da apresentação final são claros: ancoragem da PD base em série BCB/SCR, adição de LGD ao simulador de crédito, e correção da componente temperatura no score regional. Essas três mudanças elevam o rigor metodológico do tradutor financeiro de 'ordem de grandeza plausível' para 'calibração defensável'.")
add_body(doc, "O convite ao mentor e à equipe da XP é para feedback técnico específico sobre dois pontos: (1) a adequação dos multiplicadores de PD climática para o contexto real de carteiras CRA/LCA da XP; e (2) a viabilidade de integração com dados proprietários de inadimplência para calibração empírica dos parâmetros. Esses inputs transformariam o projeto de protótipo metodológico em ferramenta com potencial de uso real.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ANEXOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Anexos", level=1, size=14)

# Anexo A
add_heading(doc, "Anexo A — Fatores de Sensibilidade ENSO por UF", level=2)
add_body(doc, "Fonte: scripts/score_ufs.py — ENSO_SENS_UF. Positivo = El Niño aumenta risco; negativo = La Niña aumenta risco.", size=10)

tbl_a = doc.add_table(rows=0, cols=4)
tbl_a.style = "Table Grid"
set_table_header(tbl_a.add_row(), ["UF", "Sensibilidade", "UF", "Sensibilidade"])
enso_data = [
    ("CE","+0,95"),("RN","+0,90"),("PB","+0,90"),("PE","+0,85"),
    ("PI","+0,85"),("AL","+0,80"),("SE","+0,75"),("BA","+0,70"),
    ("RS","−0,85"),("SC","−0,75"),("PR","−0,70"),("MA","+0,65"),
    ("AM","+0,65"),("MT","+0,65"),("PA","+0,60"),("GO","+0,60"),
    ("AC","+0,60"),("TO","+0,60"),("MS","+0,55"),("DF","+0,55"),
    ("RO","+0,55"),("RR","+0,50"),("SP","+0,30"),("MG","+0,40"),
    ("ES","+0,35"),("RJ","+0,25"),("AP","+0,45"),("—","—"),
]
for i in range(0, len(enso_data), 2):
    pair = enso_data[i:i+2]
    if len(pair) == 2:
        add_table_row(tbl_a, [pair[0][0], pair[0][1], pair[1][0], pair[1][1]])
    else:
        add_table_row(tbl_a, [pair[0][0], pair[0][1], "—", "—"])

doc.add_paragraph()

# Anexo B
add_heading(doc, "Anexo B — Sensibilidade Climática por Cultura (V2 — a calibrar)", level=2)
add_body(doc, "Fatores-alvo para a V2 do Tradutor Financeiro. Valores atuais na V1: todas as culturas usam apenas o fator regional (nenhum fator cultura-específico).", size=10)

tbl_b = doc.add_table(rows=0, cols=3)
tbl_b.style = "Table Grid"
set_table_header(tbl_b.add_row(), ["Cultura", "Fator proposto (V2)", "Justificativa"])
rows_b = [
    ("Soja", "×1,0 (referência)", "Cultura de maior representatividade — base do fator 1"),
    ("Milho 1ª safra", "×1,1", "Alta sensibilidade a déficit hídrico na fase de florescimento"),
    ("Café", "×1,3", "Veranico durante floração → perda irreversível na safra seguinte"),
    ("Cana", "×0,8", "Cultura perene, maior resiliência a variações climáticas de curto prazo"),
    ("Algodão", "×1,2", "Sensível a excesso hídrico na colheita (mancha, apodrecimento)"),
    ("Arroz", "×0,9", "Irrigado no Sul (RS) tem menor exposição a déficit hídrico de sequeiro"),
    ("Feijão", "×1,1", "Ciclo curto, alta sensibilidade a estresses hídricos"),
    ("Trigo", "×1,0", "Base Sul — sensibilidade similar à soja na região"),
]
for r in rows_b:
    add_table_row(tbl_b, r)

doc.add_paragraph()
add_placeholder(doc, "Os fatores desta tabela são propostas de calibração para V2. A V1 atual não implementa fatores por cultura — todas as operações no simulador CRA/LCA recebem apenas o fator regional. Fonte de referência para calibração: EMBRAPA Agroclimatologia; relatórios técnicos de seguro rural MAPA/IRB [VERIFICAR REFERÊNCIA EXATA].")

# Anexo C
add_heading(doc, "Anexo C — Mapeamento Commodity × Região Produtora Dominante", level=2)
add_body(doc, "Fonte: scripts/app.py — _COMM_MAP. Usado nos alertas do perfil Commodities.", size=10)

tbl_c = doc.add_table(rows=0, cols=3)
tbl_c.style = "Table Grid"
set_table_header(tbl_c.add_row(), ["Região", "Commodities principais", "Contratos futuros relevantes (B3/CBOT)"])
rows_c = [
    ("Centro-Oeste", "Soja, Milho, Algodão", "SFI (soja), CCM (milho), ALF (algodão)"),
    ("Sul", "Soja, Milho, Trigo", "SFI (soja), CCM (milho), TRF (trigo)"),
    ("Nordeste", "Cana, Algodão, Milho", "ACF (açúcar cristal), ALF (algodão), CCM (milho)"),
    ("Sudeste", "Café, Cana", "ICF (café arábica), ACF (açúcar cristal)"),
    ("Norte", "Soja, Milho (expansão MATOPIBA)", "SFI (soja), CCM (milho)"),
]
for r in rows_c:
    add_table_row(tbl_c, r)

doc.add_paragraph()

# Anexo D
add_heading(doc, "Anexo D — Bibliotecas Python Utilizadas", level=2)
add_body(doc, "Conforme requirements.txt do projeto:", size=10)

libs = [
    ("requests", "HTTP — download de dados das APIs externas"),
    ("pandas", "Manipulação de DataFrames e CSVs"),
    ("numpy", "Operações numéricas vetorizadas"),
    ("streamlit", "Framework do dashboard web"),
    ("plotly / plotly.express", "Visualizações interativas (gráficos, mapas coropléticos)"),
    ("folium + streamlit-folium", "Mapa geográfico interativo com pontos de foco"),
    ("scikit-learn", "BallTree + Haversine para vizinho mais próximo municipal"),
    ("python-docx", "Geração deste relatório"),
]
tbl_d = doc.add_table(rows=0, cols=2)
tbl_d.style = "Table Grid"
set_table_header(tbl_d.add_row(), ["Biblioteca", "Uso no projeto"])
for lib, uso in libs:
    add_table_row(tbl_d, [lib, uso])

doc.add_paragraph()

# Anexo E — Glossário
add_heading(doc, "Anexo E — Glossário de Termos Técnicos", level=2)

glossario = [
    ("CRA", "Certificado de Recebíveis do Agronegócio — instrumento de securitização de crédito agro"),
    ("LCA", "Letra de Crédito do Agronegócio — título de renda fixa isento de IR vinculado ao agronegócio"),
    ("CPR", "Cédula de Produto Rural — título de crédito que representa promessa de entrega de produto rural"),
    ("PD", "Probability of Default — probabilidade de inadimplência do devedor em um horizonte de tempo"),
    ("LGD", "Loss Given Default — percentual da exposição efetivamente perdido em caso de default (1 − taxa de recuperação)"),
    ("EAD", "Exposure at Default — valor exposto no momento do default"),
    ("EL", "Expected Loss — perda esperada = EAD × PD × LGD"),
    ("HHI", "Herfindahl-Hirschman Index — medida de concentração de portfólio. HHI = Σ(share_i²)"),
    ("ONI", "Oceanic Niño Index — índice de anomalia de temperatura superficial do mar na região Niño 3.4 (Pacífico Equatorial Central)"),
    ("ENSO", "El Niño-Southern Oscillation — padrão climático de variabilidade interanual do Pacífico. El Niño: ONI ≥ +0,5°C; La Niña: ONI ≤ −0,5°C"),
    ("ERA5", "European Reanalysis — produto de reanálise climática do ECMWF/Copernicus. Grade global de 31km, disponível via Open-Meteo sem autenticação"),
    ("IFRS9", "International Financial Reporting Standard 9 — norma contábil que exige provisionamento prospectivo de perdas esperadas em crédito"),
    ("Basel III", "Acordo de Basileia III — framework regulatório internacional de adequação de capital para bancos"),
    ("TCFD", "Task Force on Climate-related Financial Disclosures — framework de divulgação de riscos climáticos para empresas e investidores"),
    ("NGFS", "Network for Greening the Financial System — rede de bancos centrais e reguladores para integração de risco climático em políticas financeiras"),
    ("FIAGRO", "Fundo de Investimento nas Cadeias Produtivas Agroindustriais — veículo de investimento brasileiro em CRA, CPR e ativos agrícolas"),
]

tbl_e = doc.add_table(rows=0, cols=2)
tbl_e.style = "Table Grid"
set_table_header(tbl_e.add_row(), ["Termo", "Definição"])
for termo, defn in glossario:
    add_table_row(tbl_e, [termo, defn], bold_first=True)

# ── Salva ─────────────────────────────────────────────────────────────────────
out_path = "/Users/joao/climacredit/docs/ClimaCredit_Relatorio_Tecnico_v1.docx"
doc.save(out_path)
print(f"Salvo: {out_path}")
